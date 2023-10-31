from flask import Flask, request, send_file, make_response, render_template, url_for, markup
import json
from docx import Document
import os
from io import BytesIO
import markdown2



app = Flask(__name__)
# markdown2.markdown(app)

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/page1')
def page1():
    return render_template('page1.html')

@app.route('/page2')
def page2():
    return render_template('page2.html')

# This method takes the uploaded file and stores it on the server in the uploads folder.
# This method also outputs where the file was saved.

# Define a directory to store uploaded files
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER



# @app.route('/upload1', methods=['POST'])
# def upload_file_1():
#     # Function implementation
#     return 'Upload Method 1'
@app.route('/upload1', methods=['GET', 'POST'])
def upload_file_1():
    if request.method == 'POST':
        file = request.files['file']
        if file and file.filename.endswith('.json'):
            json_data = json.load(file)
            markdown_text = "# Table of JSON Data\n\n"
            markdown_text += "| Key | Value |\n"
            markdown_text += "|---|---|\n"
            for key, value in json_data.items():
                markdown_text += f"| {key} | {value} |\n"
            
            output_filename = 'output1.md'
            output_filepath = os.path.join(UPLOAD_FOLDER, output_filename)
            with open(output_filepath, 'w') as md_file:
                md_file.write(markdown_text)

            return render_template('display_markdown.html', markdown_content=markup(markdown_text))

    return render_template('upload.html')
# @app.route('/upload1', methods=['GET', 'POST'])
# def upload_file_1():
#     if request.method == 'POST':
#         if 'file' not in request.files:
#             return 'No file part'
        
#         file = request.files['file']
#         if file.filename == '':
#             return 'No selected file'

#         if file and file.filename.endswith('.json'):
#             try:
#                 # Load JSON data from the uploaded file
#                 json_data = json.load(file)
                
#                 # Generate the markdown table
#                 markdown_text = "# Table of JSON Data\n\n"
#                 markdown_text += "| Key | Value |\n"
#                 markdown_text += "|---|---|\n"
#                 for key, value in json_data.items():
#                     markdown_text += f"| {key} | {value} |\n"
                
#                 # Save the markdown text to a file
#                 output_filename = 'output.md'
#                 output_filepath = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
#                 with open(output_filepath, 'w') as md_file:
#                     md_file.write(markdown_text)


#                 return render_template('display_markdown.html', markdown_content=markdown_text)


#                 # # Return the path to the user
#                 # return f"The markdown file has been saved to: {output_filepath}"

#             except Exception as e:
#                 return str(e)
            
#     return render_template('index.html')
#     # return 'Invalid file format. Please upload a JSON file.'


@app.route('/upload2', methods=['POST'])
def upload_file_2():
    if 'file' not in request.files:
        return 'No file part'
    
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'

    if file and file.filename.endswith('.json'):
        try:
            json_data = json.load(file)
            markdown_text = "# Table of JSON Data\n\n"
            markdown_text += "| Key | Value |\n"
            markdown_text += "|---|---|\n"

            for key, value in json_data.items():
                markdown_text += f"| {key} | {value} |\n"

            response = make_response(markdown_text)
            response.headers["Content-Disposition"] = "attachment; filename=output.md"
            response.mimetype = 'text/markdown'
            return response

        except Exception as e:
            return str(e)
    
    return 'Invalid file format. Please upload a JSON file.'

@app.route('/upload3', methods=['POST'])
def upload_file_3():
    if 'file' not in request.files:
        return 'No file part'
    
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'

    if file and file.filename.endswith('.json'):
        try:
            json_data = json.load(file)
            doc = Document()
            for key, value in json_data.items():
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Key'
                hdr_cells[1].text = 'Value'
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
                doc.add_paragraph()

            # Save the document to a BytesIO object
            f = BytesIO()
            doc.save(f)
            f.seek(0)

            return send_file(f, as_attachment=True, download_name='output.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        except Exception as e:
            return str(e)
    
    return 'Invalid file format. Please upload a JSON file.'

if __name__ == '__main__':
    app.run(debug=True)
